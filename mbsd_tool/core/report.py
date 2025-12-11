from __future__ import annotations

import json
from pathlib import Path
from typing import TextIO

from mbsd_tool.core.models import ScanResult, VulnerabilityFinding
from datetime import datetime
from jinja2 import Template


def export_markdown(result: ScanResult, path: str | Path) -> None:
    p = Path(path)
    with p.open("w", encoding="utf-8") as f:
        _write_markdown(result, f)


def _write_markdown(result: ScanResult, f: TextIO) -> None:
    f.write(f"# 脆弱性診断レポート\n\n")
    f.write(f"対象: {result.target}\n\n")
    f.write(f"モード: {result.mode.value}\n\n")
    for endpoint, vulns in result.vulns_by_endpoint.items():
        f.write(f"## {endpoint}\n\n")
        if not vulns:
            f.write("- 問題は検出されませんでした\n\n")
            continue
        for v in vulns:
            f.write(f"- 脆弱性: {v.name}\n")
            f.write(f"  - 重要度: {v.severity}\n")
            if v.evidence:
                f.write(f"  - 証拠: {v.evidence}\n")
            if v.reproduction_steps:
                f.write(f"  - 再現手順:\n")
                for s in v.reproduction_steps:
                    f.write(f"    1. {s}\n")
            f.write("\n")


def export_json(result: ScanResult, path: str | Path) -> None:
    p = Path(path)
    with p.open("w", encoding="utf-8") as f:
        json.dump(result.model_dump(), f, indent=2, ensure_ascii=False)


def _filter_fields(v: VulnerabilityFinding, fields: dict) -> dict:
    d = {}
    if fields.get("name"):
        d["name"] = v.name
    if fields.get("severity"):
        d["severity"] = v.severity
    if fields.get("evidence") and v.evidence:
        d["evidence"] = v.evidence
    if fields.get("repro") and v.reproduction_steps:
        d["reproduction_steps"] = v.reproduction_steps
    if fields.get("notes") and v.notes:
        d["notes"] = v.notes
    return d


def export_html(result: ScanResult, path: str | Path, fields: dict | None = None) -> None:
    # Block style with per-vulnerability section, red large title, and a blank line between items
    tpl = Template(
        """
        <!doctype html><html lang="ja"><head>
        <meta charset="utf-8"/>
        <title>脆弱性診断レポート</title>
        <style>
        body{font-family:-apple-system,BlinkMacSystemFont,Segoe UI, sans-serif; padding:16px;}
        h1{margin:0 0 10px}
        h2{margin:0 0 6px; color:#b91c1c; font-size:22px}
        .block{margin:0 0 16px; padding:12px; border:1px solid #e5e7eb; border-radius:6px; page-break-after: always; break-after: page;}
        .label{font-weight:600; color:#374151; margin-top:6px}
        .mono{font-family:ui-monospace, SFMono-Regular, Menlo, monospace}
        .box{border:1px solid #e5e7eb; padding:8px; background:#fafafa; border-radius:4px; margin-bottom:8px}
        .sev{display:inline-block; padding:2px 8px; border-radius:999px; font-size:12px;}
        .sev-high{background:#7f1d1d; color:#fff}
        .sev-med{background:#7c2d12; color:#fff}
        .sev-low{background:#064e3b; color:#fff}
        .sev-info{background:#1e3a8a; color:#fff}
        .meta{color:#6b7280; margin-bottom:8px}
        </style>
        </head><body>
        <h1>脆弱性診断レポート</h1>
        <div class="meta">対象: {{ target }} / モード: {{ mode }}</div>
        {% set i = 1 %}
        {% for endpoint, vulns in items %}
          {% for v in vulns %}
            <div class="block">
              <h2>[{{ i }}] {{ v.name }} <span class="sev {{ 'sev-high' if v.severity=='高' else ('sev-med' if v.severity=='中' else ('sev-low' if v.severity=='低' else 'sev-info')) }}">{{ v.severity }}</span></h2>
              <div class="label">対象</div>
              <div class="box">{{ v.feature_name or '不明' }}<br/><span class="mono">{{ endpoint }}</span></div>
              {% if v.explanation %}
              <div class="label">解説</div>
              <div class="box">{{ v.explanation }}</div>
              {% endif %}
              {% if v.impact %}
              <div class="label">想定される被害・影響</div>
              <div class="box">{{ v.impact }}</div>
              {% endif %}
              {% if v.remediation %}
              <div class="label">対策</div>
              <div class="box">{{ v.remediation }}</div>
              {% endif %}
              {% if v.evidence %}
              <div class="label">証拠</div>
              <div class="box">{{ v.evidence }}</div>
              {% endif %}
              {% if v.reproduction_steps %}
              <div class="label">再現手順</div>
              <div class="box"><ol>{% for s in v.reproduction_steps %}<li>{{ s }}</li>{% endfor %}</ol></div>
              {% endif %}
              {% if v.notes %}
              <div class="label">備考</div>
              <div class="box">{{ v.notes }}</div>
              {% endif %}
            </div>
            {% set i = i + 1 %}
          {% endfor %}
        {% endfor %}
        </body></html>
        """
    )
    html = tpl.render(target=result.target, mode=result.mode.value, items=result.vulns_by_endpoint.items())
    Path(path).write_text(html, encoding="utf-8")


def export_pdf(result: ScanResult, path: str | Path, fields: dict | None = None) -> None:
    # Use the AIEL-style PDF layout as the default to match the provided resources
    export_aiel_pdf(result, path)


def export_report(result: ScanResult, path: str | Path, fmt: str, fields: dict | None = None) -> None:
    fmt = fmt.lower()
    if fmt == "markdown":
        export_markdown(result, path)
    elif fmt == "json":
        export_json(result, path)
    elif fmt == "html":
        export_html(result, path, fields)
    elif fmt == "pdf":
        export_pdf(result, path, fields)
    elif fmt in ("pdf (aiel)", "aiel_pdf", "pdf_aiel"):
        export_aiel_pdf(result, path)
    elif fmt in ("docx (aiel)", "aiel_docx", "docx_aiel"):
        export_aiel_docx(result, path)
    elif fmt in ("指定形式(markdown)", "company_markdown", "custom_markdown"):
        export_company_markdown(result, path)
    elif fmt in ("指定形式(html)", "company_html", "custom_html"):
        export_company_html(result, path)
    elif fmt in ("指定形式(pdf)", "company_pdf", "custom_pdf"):
        export_company_pdf(result, path)
    else:
        raise ValueError(f"未知の形式: {fmt}")


def _severity_counts(result: ScanResult) -> dict:
    sev_counts: dict[str, int] = {"高": 0, "中": 0, "低": 0, "情報": 0}
    for _, vulns in result.vulns_by_endpoint.items():
        for v in vulns:
            sev_counts[v.severity] = sev_counts.get(v.severity, 0) + 1
    return sev_counts

def _guess_request_name(url: str) -> str:
    try:
        from urllib.parse import urlparse
        pr = urlparse(url)
        path = pr.path.rstrip('/') or '/'
        if path == '/':
            return 'トップページ'
        last = path.split('/')[-1]
        mapping = {
            'login': 'ログイン', 'logout': 'ログアウト', 'top': 'トップページ', 'edit': '編集',
            'complete': '完了', 'list': '一覧', 'detail': '詳細', 'password': 'パスワード',
        }
        return mapping.get(last.lower(), last)
    except Exception:
        return 'リクエスト'


def export_aiel_pdf(result: ScanResult, path: str | Path) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.lib.colors import Color
    except Exception as e:
        raise RuntimeError("reportlabがインストールされていません。'pip install reportlab' を実行してください。") from e

    # フォント
    try:
        pdfmetrics.registerFont(UnicodeCIDFont('HeiseiKakuGo-W5'))
        JP_FONT = 'HeiseiKakuGo-W5'
        JP_FONT_BOLD = 'HeiseiKakuGo-W5'
    except Exception as e:
        raise RuntimeError("PDF生成に必要な日本語フォント（平成角ゴシック）が見つかりません。") from e

    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    margin = 18 * mm
    x0 = margin
    y = height - margin

    # カバー
    c.setFillColorRGB(0.15, 0.18, 0.22)
    c.rect(0, height*0.65, width, height*0.35, fill=1, stroke=0)
    c.setFillColorRGB(1, 1, 1)
    c.setFont(JP_FONT_BOLD, 22)
    c.drawString(x0, height*0.77, "脆弱性診断レポート")
    c.setFont(JP_FONT, 12)
    c.drawString(x0, height*0.77 - 24, f"対象: {result.target}")
    c.drawString(x0, height*0.77 - 42, f"モード: {result.mode.value}")
    c.drawString(x0, height*0.77 - 60, f"作成日: {datetime.now().strftime('%Y-%m-%d')}")

    # 要約（重要度別件数）
    y = height*0.62
    c.setFillColorRGB(0, 0, 0)
    c.setFont(JP_FONT_BOLD, 14)
    c.drawString(x0, y, "サマリー")
    y -= 10
    sev = _severity_counts(result)
    # 色（落ち着いたトーン）
    colors = {
        '高': Color(0.50, 0.12, 0.12),
        '中': Color(0.60, 0.32, 0.10),
        '低': Color(0.12, 0.33, 0.18),
        '情報': Color(0.12, 0.23, 0.54),
    }
    box_w = (width - 2*margin - 24) / 4.0
    box_h = 32
    y -= (box_h + 6)
    i = 0
    for k in ("高", "中", "低", "情報"):
        cx = x0 + i * (box_w + 8)
        c.setFillColor(colors[k])
        c.roundRect(cx, y, box_w, box_h, 4, fill=1, stroke=0)
        c.setFillColorRGB(1, 1, 1)
        c.setFont(JP_FONT_BOLD, 12)
        c.drawString(cx + 8, y + box_h - 12, k)
        c.setFont(JP_FONT_BOLD, 16)
        c.drawRightString(cx + box_w - 8, y + 10, str(sev.get(k, 0)))
        i += 1

    # 概要テキスト
    y -= 46
    c.setFillColorRGB(0, 0, 0)
    c.setFont(JP_FONT, 10)
    c.drawString(x0, y, f"エンドポイント数: {len(result.endpoints)} / 総検出数: {sum(sev.values())}")
    y -= 24

    # 詳細セクション
    def write_line(text: str, leading: float = 12.0, indent: float = 0.0, bold: bool = False, size: float | None = None):
        nonlocal y
        if y < margin:
            c.showPage(); y = height - margin
            c.setFont(JP_FONT, 10)
        c.setFont(JP_FONT_BOLD if bold else JP_FONT, (size if size else (11 if bold else 10)))
        c.drawString(x0 + indent, y, text)
        y -= leading

    c.setFont(JP_FONT_BOLD, 14)
    write_line("詳細", 18, 0, True)
    idx = 1
    for endpoint, vulns in result.vulns_by_endpoint.items():
        if not vulns:
            continue
        for v in vulns:
            # 各脆弱性は改ページで区切る
            write_line(endpoint, 16, 0, True)
            c.setFillColorRGB(0.73, 0.0, 0.0)
            write_line(f"[{idx}] {v.name} / 重要度: {v.severity}", 16, 6, True, size=13)
            c.setFillColorRGB(0, 0, 0)
            if v.explanation:
                write_line("解説:", 12, 10, True)
                for line in (v.explanation or '').splitlines() or ['']:
                    write_line(line, 12, 14)
                write_line("", 8)
            if v.evidence:
                write_line("証拠:", 12, 10, True)
                write_line(v.evidence, 12, 14)
                write_line("", 8)
            if v.reproduction_steps:
                write_line("再現手順:", 12, 10, True)
                for s in v.reproduction_steps:
                    write_line(f"- {s}", 12, 14)
                write_line("", 8)
            if v.impact:
                write_line("影響:", 12, 10, True)
                write_line(v.impact, 12, 14)
                write_line("", 8)
            if v.remediation:
                write_line("対策:", 12, 10, True)
                write_line(v.remediation, 12, 14)
                write_line("", 8)
            if v.notes:
                write_line("備考:", 12, 10, True)
                write_line(v.notes, 12, 14)
                write_line("", 8)
            idx += 1
            # 改ページ
            c.showPage(); y = height - margin
            c.setFont(JP_FONT, 10)

    c.save()


def export_aiel_docx(result: ScanResult, path: str | Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as e:
        raise RuntimeError("python-docx がインストールされていません") from e

    doc = Document()
    title = doc.add_heading('脆弱性診断レポート', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph(f"対象: {result.target}\nモード: {result.mode.value}\n作成日: {datetime.now().strftime('%Y-%m-%d')}")
    p.style.font.size = Pt(11)

    doc.add_heading('サマリー', level=1)
    sev = _severity_counts(result)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = f"高: {sev.get('高',0)}"; hdr[1].text = f"中: {sev.get('中',0)}"; hdr[2].text = f"低: {sev.get('低',0)}"; hdr[3].text = f"情報: {sev.get('情報',0)}"

    doc.add_heading('詳細', level=1)
    idx = 1
    for endpoint, vulns in result.vulns_by_endpoint.items():
        if not vulns:
            continue
        doc.add_heading(endpoint, level=2)
        for v in vulns:
            p_title = doc.add_heading(f"[{idx}] {v.name} / 重要度: {v.severity}", level=3)
            # タイトルを赤色・やや大きめに
            for run in p_title.runs:
                run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
                run.font.size = Pt(14)
            if v.explanation:
                doc.add_paragraph("解説:")
                doc.add_paragraph(v.explanation)
                doc.add_paragraph("")
            if v.evidence:
                doc.add_paragraph("証拠:")
                doc.add_paragraph(v.evidence)
                doc.add_paragraph("")
            if v.reproduction_steps:
                doc.add_paragraph("再現手順:")
                for s in v.reproduction_steps:
                    doc.add_paragraph(s, style='List Bullet')
                doc.add_paragraph("")
            if v.impact:
                doc.add_paragraph("影響:")
                doc.add_paragraph(v.impact)
                doc.add_paragraph("")
            if v.remediation:
                doc.add_paragraph("対策:")
                doc.add_paragraph(v.remediation)
                doc.add_paragraph("")
            if v.notes:
                doc.add_paragraph("備考:")
                doc.add_paragraph(v.notes)
            # 1件ごとに改ページ
            doc.add_page_break()
            idx += 1

    doc.save(str(path))


def export_company_markdown(result: ScanResult, path: str | Path) -> None:
    p = Path(path)
    with p.open("w", encoding="utf-8") as f:
        idx = 1
        for endpoint, vulns in result.vulns_by_endpoint.items():
            for v in vulns:
                f.write(f"[{idx}]\t{v.name}\n")
                f.write("\uf0a0\t対象\n")
                f.write(f"{v.feature_name or '機能名'}\n")
                f.write(f"{endpoint}\n\n")
                f.write("\uf0a0\t危険度\n")
                sev_map = {"高": "High", "中": "Medium", "低": "Low", "情報": "Info"}
                f.write(f"{sev_map.get(v.severity, v.severity)}\n\n")
                f.write("\uf0a0\t解説\n")
                f.write(f"{v.explanation or v.evidence or ''}\n\n")
                f.write("\uf0a0\t想定される被害・影響\n")
                f.write(f"{v.impact or ''}\n\n")
                f.write("\uf0a0\t対策\n")
                f.write(f"{v.remediation or ''}\n\n")
                f.write("\uf0a0\t備考\n")
                f.write(f"{v.notes or ''}\n\n")
                idx += 1


def export_company_html(result: ScanResult, path: str | Path) -> None:
    flat = []
    for endpoint, vulns in result.vulns_by_endpoint.items():
        for v in vulns:
            flat.append((endpoint, v))
    paths_data = [{"no": i, "name": _guess_request_name(u), "url": u} for i, u in enumerate(result.endpoints, start=1)]
    tpl = Template(
        """
        <!doctype html><html lang="ja"><head>
        <meta charset="utf-8"/>
        <title>脆弱性診断レポート（指定形式）</title>
        <style>
        :root{--pad:16px; --title:#b91c1c; --uline:3px}
        h1{margin:0 0 10px}
        .title-wrap{margin-left:calc(var(--pad) * -1); margin-right:calc(var(--pad) * -1); padding:0 var(--pad); border-bottom:var(--uline) solid var(--title);}
        .title{color:var(--title); font-size:22px; margin:0 0 8px}
        body{font-family:-apple-system,BlinkMacSystemFont,Segoe UI, sans-serif; padding:16px;}
        h2{margin:0 0 6px; color:#b91c1c; font-size:22px}
        .block{margin:0 0 16px; page-break-after: always; break-after: page;}
        .label{font-weight:600;}
        .mono{font-family:ui-monospace, SFMono-Regular, Menlo, monospace}
        .content{margin-bottom:8px}
        table{border-collapse:collapse; width:100%; margin:8px 0}
        th,td{border:1px solid #ddd; padding:6px; text-align:left}
        th{background:#f5f5f5}
        .page{page-break-after: always; break-after: page;}
        </style>
        </head><body>
        <h1>脆弱性診断レポート（指定形式）</h1>
        <div class="page">
          <div class="title-wrap"><h2 class="title">サマリー</h2></div>
          <table>
            <thead><tr><th>No</th><th>脆弱性</th><th>危険度</th><th>パス</th></tr></thead>
            <tbody>
            {% set sevmap = {'高':'High','中':'Medium','低':'Low','情報':'Info'} %}
            {% for i, item in enumerate(flat, start=1) %}
              {% set endpoint, v = item %}
              <tr><td>[{{ i }}]</td><td>{{ v.name }}</td><td>{{ sevmap.get(v.severity, v.severity) }}</td><td><span class="mono">{{ endpoint }}</span></td></tr>
            {% endfor %}
            </tbody>
          </table>
        </div>
        <div class="page">
          <div class="title-wrap"><h2 class="title">パス一覧</h2></div>
          <table>
            <thead><tr><th>No</th><th>リクエスト名</th><th>URL</th></tr></thead>
            <tbody>
            {% for it in paths_data %}
              <tr><td>[{{ it.no }}]</td><td>{{ it.name }}</td><td><span class="mono">{{ it.url }}</span></td></tr>
            {% endfor %}
            </tbody>
          </table>
        </div>
        {% set i = 1 %}
        {% for endpoint, vulns in items %}
          {% for v in vulns %}
            <div class="block">
              <div class="title-wrap"><h2 class="title">[{{ i }}]	{{ v.name }}</h2></div>
              <div class="label">対象</div>
              <div class="content">{{ v.feature_name or '機能名' }}<br/><span class="mono">{{ endpoint }}</span></div>
              <div class="label">危険度</div>
              {% set sevmap = {'高':'High','中':'Medium','低':'Low','情報':'Info'} %}
              <div class="content">{{ sevmap.get(v.severity, v.severity) }}</div>
              <div class="label">解説</div>
              <div class="content">{{ v.explanation or v.evidence or '' }}</div>
              <div class="label">想定される被害・影響</div>
              <div class="content">{{ v.impact or '' }}</div>
              <div class="label">対策</div>
              <div class="content">{{ v.remediation or '' }}</div>
              <div class="label">備考</div>
              <div class="content">{{ v.notes or '' }}</div>
            </div>
            {% set i = i + 1 %}
          {% endfor %}
        {% endfor %}
        </body></html>
        """
    )
    html = tpl.render(items=result.vulns_by_endpoint.items(), flat=flat, paths_data=paths_data)
    Path(path).write_text(html, encoding="utf-8")


def export_company_pdf(result: ScanResult, path: str | Path) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    except Exception as e:
        raise RuntimeError("reportlabがインストールされていません。'pip install reportlab' を実行してください。") from e
    try:
        pdfmetrics.registerFont(UnicodeCIDFont('HeiseiKakuGo-W5'))
        JP_FONT = 'HeiseiKakuGo-W5'
        JP_FONT_BOLD = 'HeiseiKakuGo-W5'
    except Exception as e:
        raise RuntimeError("PDF生成に必要な日本語フォント（平成角ゴシック）が見つかりません。") from e
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    margin = 15 * mm
    x = margin
    y = height - margin
    def write_line(text: str, leading: float = 12.0, indent: float = 0.0, bold: bool = False, size: float | None = None):
        nonlocal y
        if y < margin:
            c.showPage(); y = height - margin
            c.setFont(JP_FONT_BOLD if bold else JP_FONT, 10 if not bold else 11)
        c.setFont(JP_FONT_BOLD if bold else JP_FONT, (size if size else (10 if not bold else 11)))
        c.drawString(x + indent, y, text)
        y -= leading

    # Flatten for stable numbering
    flat = []
    for endpoint, vulns in result.vulns_by_endpoint.items():
        for v in vulns:
            flat.append((endpoint, v))

    # サマリー
    c.setFont(JP_FONT_BOLD, 14); write_line("脆弱性診断レポート（指定形式）", 18)
    write_line("サマリー", 16, 0, True)
    sev_map = {"高": "High", "中": "Medium", "低": "Low", "情報": "Info"}
    idx = 1
    for endpoint, v in flat:
        write_line(f"[{idx}] {v.name} / {sev_map.get(v.severity, v.severity)}", 12, 0, False)
        write_line(endpoint, 12, 10)
        idx += 1
    # 改ページ
    c.showPage(); y = height - margin

    # パス一覧
    c.setFont(JP_FONT_BOLD, 14); write_line("脆弱性診断レポート（指定形式）", 18)
    write_line("パス一覧", 16, 0, True)
    paths_data = [{"no": i, "name": _guess_request_name(u), "url": u} for i, u in enumerate(result.endpoints, start=1)]
    for it in paths_data:
        write_line(f"[{it['no']}] {it['name']}", 12, 0)
        write_line(it['url'], 12, 10)
    c.showPage(); y = height - margin

    # 詳細（1件1ページ）
    c.setFont(JP_FONT_BOLD, 14); write_line("脆弱性診断レポート（指定形式）", 18)
    idx = 1
    for endpoint, v in flat:
        # タイトルを大きめ＆赤色、下線をページ端から端へ
        c.setFillColorRGB(0.73, 0.0, 0.0)
        write_line(f"[{idx}] {v.name}", 16, 0, True, size=13)
        c.setStrokeColorRGB(0.73, 0.0, 0.0); c.setLineWidth(2.0)
        c.line(0, y + 12, width, y + 12)
        c.setFillColorRGB(0, 0, 0)
        write_line("対象", 12, 0, True)
        write_line(f"{v.feature_name or '不明'}", 12, 10)
        write_line(endpoint, 12, 10)
        write_line("", 8)
        write_line("危険度", 12, 0, True)
        write_line(sev_map.get(v.severity, v.severity), 12, 10)
        write_line("", 8)
        write_line("解説", 12, 0, True)
        write_line(v.explanation or v.evidence or "", 12, 10)
        write_line("", 8)
        write_line("想定される被害・影響", 12, 0, True)
        write_line(v.impact or "", 12, 10)
        write_line("", 8)
        write_line("対策", 12, 0, True)
        write_line(v.remediation or "", 12, 10)
        write_line("", 8)
        write_line("備考", 12, 0, True)
        write_line(v.notes or "", 12, 10)
        write_line("", 8)
        idx += 1
        c.showPage(); y = height - margin
        c.setFont(JP_FONT_BOLD, 14); write_line("脆弱性診断レポート（指定形式）", 18)
    c.save()
